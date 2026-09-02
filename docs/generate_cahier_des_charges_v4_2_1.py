from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT = Path(__file__).with_name("AlertContacts_Cahier_des_charges_V4_2_1.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(25, 25, 25)
MUTED = RGBColor(95, 95, 95)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def set_run(run, size=None, bold=None, italic=None, color=None):
    font = run.font
    font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        font.size = Pt(size)
    if bold is not None:
        font.bold = bold
    if italic is not None:
        font.italic = italic
    if color is not None:
        font.color.rgb = color


def paragraph(text="", style=None, bold=False, italic=False, size=None, color=None, after=6, before=0):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if text:
        r = p.add_run(text)
        set_run(r, size=size, bold=bold, italic=italic, color=color or INK)
    return p


def h1(text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    return p.add_run(text)


def h2(text):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p.add_run(text)


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    p.add_run(text)
    return p


def number(text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    p.add_run(text)
    return p


def table(headers, rows, widths, header_fill=LIGHT_GRAY):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    tr_pr = t.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    hdr = t.rows[0].cells
    for i, text in enumerate(headers):
        set_cell_shading(hdr[i], header_fill)
        run = hdr[i].paragraphs[0].add_run(text)
        set_run(run, bold=True, color=INK)
    for row in rows:
        cells = t.add_row().cells
        for i, text in enumerate(row):
            cells[i].paragraphs[0].add_run(str(text))
    set_table_widths(t, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def callout(title, body):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    tr_pr = t.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    cell = t.rows[0].cells[0]
    set_cell_shading(cell, LIGHT_BLUE)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_run(r, bold=True, color=DARK_BLUE)
    p.add_run("\n")
    r2 = p.add_run(body)
    set_run(r2, color=INK)
    set_table_widths(t, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_metadata(label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label + " ")
    set_run(r, bold=True, color=INK)
    r = p.add_run(value)
    set_run(r, color=INK)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(11)
styles["Normal"].font.color.rgb = INK
styles["Normal"].paragraph_format.space_after = Pt(6)
styles["Normal"].paragraph_format.line_spacing = 1.10

for name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
]:
    style = styles[name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = True
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = footer.add_run("AlertContacts - Cahier des charges V4.2.1")
set_run(r, size=9, color=MUTED)

paragraph("CAHIER DES CHARGES", bold=True, size=23, color=INK, after=4, before=16)
paragraph("AlertContacts - version fonctionnelle mise a jour", size=15, color=MUTED, after=16)
add_metadata("Destinataire :", "Cliente AlertContacts")
add_metadata("Document :", "Specification fonctionnelle et technique de l'application mobile et du backend")
add_metadata("Version produit observee :", "Flutter 4.2.1+50, backend Laravel associe")
add_metadata("Date :", date.today().strftime("%d/%m/%Y"))
add_metadata("Statut :", "Document de reference mis a jour avec les nouvelles fonctionnalites developpees")
paragraph("")

callout(
    "Conclusion",
    "AlertContacts est aujourd'hui une application de securite personnelle et familiale structuree autour de quatre espaces : Carte, Proches, Alertes et Traceurs. Les ajouts majeurs depuis le precedent cahier des charges sont le socle V4/V4.1, le module Trajets, RevenueCat, le mode invisible Premium et la nouvelle feature autonome de traceurs GPS reels, avec une partie gratuite limitee et des fonctions de suivi actif reservees au Premium.",
)

h1("1. Objectif du produit")
paragraph(
    "AlertContacts permet a un utilisateur de rassurer et proteger ses proches au moyen du partage de position consenti, de zones de securite, d'alertes en temps reel, de signalements communautaires et, desormais, de traceurs GPS physiques."
)
paragraph(
    "Le produit ne doit jamais etre presente comme une garantie de protection absolue. Il fournit des informations, des alertes et des outils d'aide a la decision, dans un cadre ou le consentement et la confidentialite restent centraux."
)

h1("2. Perimetre fonctionnel actuel")
table(
    ["Module", "Fonctions couvertes", "Etat"],
    [
        ["Carte", "Carte Google Maps, position utilisateur, proches, zones, alertes, trajets, etats de connectivite.", "Developpe"],
        ["Proches", "Invitation, acceptation, relation bidirectionnelle, statut, permissions de partage, suppression.", "Developpe"],
        ["Zones de securite", "Creation, modification, suppression, assignation a des proches, alertes entree/sortie.", "Developpe"],
        ["Alertes", "Flux d'alertes, signalements communautaires, confirmation, resolution, abus, lecture/non-lu.", "Developpe"],
        ["Trajets", "Recherche, itineraire, incidents sur trajet, contournement, historique API, destinations recentes.", "Developpe"],
        ["Traceurs GPS", "Ajout/gestion d'un traceur, telemetry serveur, derniere position, Premium pour suivi actif/historique/zones.", "Nouveau"],
        ["Abonnement", "RevenueCat, entitlement Premium, paywall, webhook, suspension des traceurs actifs en expiration.", "Developpe"],
        ["Parametres", "Profil, notifications, confidentialite, aide, feedback, export et suppression de compte.", "Developpe"],
    ],
    [1700, 5860, 1800],
)

h1("3. Parcours utilisateur")
h2("3.1 Onboarding et authentification")
bullet("Splash screen V4, onboarding pre-authentification et personnalisation de l'usage.")
bullet("Authentification via Firebase : Google, Apple et magic link email selon la plateforme et la configuration.")
bullet("Echange du token Firebase contre un token API Laravel/Sanctum.")
bullet("Gestion des deep links pour les invitations et les liens d'authentification.")
bullet("Demandes de permissions faites en contexte : localisation, localisation en arriere-plan et notifications.")

h2("3.2 Navigation principale")
paragraph("L'application dispose maintenant de quatre onglets principaux : Carte, Proches, Alertes et Traceurs.")
bullet("Carte : vision centrale de la position, des proches, des zones, des alertes et des trajets.")
bullet("Proches : gestion des relations, invitations, statuts et permissions.")
bullet("Alertes : consultation des alertes proches, zones et communaute, avec filtre et badge non-lu.")
bullet("Traceurs : nouvel espace independant pour les traceurs GPS physiques.")

h1("4. Module Carte")
bullet("Affichage Google Maps avec recentrage sur la position courante.")
bullet("Publication de la position du mobile dans Firebase Realtime Database pour l'affichage temps reel par les proches autorises.")
bullet("Envoi batch des positions vers Laravel pour le geoprocessing et l'historique.")
bullet("Affichage des proches avec statut actif, hors ligne, partage non reciproque ou position en pause.")
bullet("Affichage des zones de securite et des zones/alertes de danger sur le viewport courant.")
bullet("Cache de viewport pour limiter les rechargements inutiles.")
bullet("Bannieres d'etat : hors ligne, batterie faible d'un proche, GPS imprecis, mode invisible actif.")
bullet("Acces aux zones via le bouton de calques et aux parametres via l'avatar.")

h1("5. Module Proches")
paragraph("Le module Proches gere les utilisateurs AlertContacts relies entre eux. Il ne doit pas etre confondu avec les traceurs GPS physiques, qui disposent de leur propre module.")
bullet("Ajout d'un proche par invitation partageable et deep link.")
bullet("Acceptation ou refus d'une invitation.")
bullet("Relation bidirectionnelle : un utilisateur peut voir un proche uniquement si le partage est autorise dans le sens attendu.")
bullet("Gestion des etats : connecte, en attente, hors ligne, position en pause, partage non visible.")
bullet("Permissions de partage prevues : position, batterie, evenements de zone et vitesse.")
bullet("Suppression d'un proche avec retrait de la relation dans les deux sens.")
bullet("Limite gratuite actuelle : 1 proche. Au-dela, l'app presente le paywall Premium et le serveur conserve la decision finale.")

h1("6. Zones de securite")
bullet("Creation de zones depuis l'interface Carte/Zones, avec choix du centre, du rayon, du nom et de l'icone.")
bullet("Rayon configurable selon la configuration actuelle : minimum 50 m, maximum 500 m, valeur par defaut 150 m.")
bullet("Le backend sait stocker des zones circulaires et polygonales ; l'interface actuelle privilegie le parcours circulaire.")
bullet("Assignation des zones uniquement a des proches acceptes.")
bullet("Notifications d'entree et de sortie, avec parametrage notify_entry et notify_exit.")
bullet("Suppression de zone et dissociation des assignations.")
bullet("Limite gratuite actuelle : 1 zone de securite. Les zones illimitees sont reservees au Premium.")

h1("7. Alertes communautaires et incidents")
paragraph("Le socle V4.1 separe le signalement brut envoye par un utilisateur et l'incident consolide affiche dans l'application.")
bullet("Creation gratuite de signalements communautaires afin d'augmenter la densite de contribution.")
bullet("Taxonomie d'incidents avec severite, description, precision GPS et trace GPS optionnelle.")
bullet("Detection de doublons avant creation pour transformer un doublon potentiel en confirmation utile.")
bullet("Clustering spatio-temporel des signalements compatibles pour produire un incident unique.")
bullet("Consultation des incidents actifs dans le viewport courant.")
bullet("Actions utilisateur : confirmer, indiquer que c'est termine, signaler un abus.")
bullet("Flux Alertes avec filtres Tout, Proches, Zones et Communaute, groupement par date et marquage lu/non-lu.")
bullet("Historique local des evenements recents, avec coupure gratuite a 24 h selon la configuration.")

h1("8. Module Trajets")
paragraph("Le module Trajets ajoute une logique de navigation prudente sans promettre un trajet parfaitement securise.")
bullet("Recherche d'itineraire avec point de depart, destination, autocompletion et inversion depart/arrivee.")
bullet("Modes de transport prevus dans l'interface : voiture, marche, deux-roues.")
bullet("Calcul d'itineraire cote serveur via le service de routage configure.")
bullet("Detection d'incidents sur le trajet et affichage d'un avertissement avant le depart.")
bullet("Possibilite de demander un contournement lorsqu'un incident routable affecte l'itineraire.")
bullet("Contournement gratuit et illimite pour les incidents de severite elevee ; quota gratuit mensuel pour les incidents faibles/moyens.")
bullet("Historique des trajets cote API, destinations recentes et transitions start, end, cancel.")
bullet("Surveillance pendant trajet reservee aux utilisateurs Premium selon la configuration backend.")

h1("9. Nouvelle feature : traceurs GPS physiques")
callout(
    "Decision produit",
    "La feature Traceurs est independante de l'ajout d'un proche AlertContacts. Elle sert aux cas ou la personne suivie ne dispose pas de telephone, mais porte un traceur GPS reel.",
)
bullet("Un traceur possede au minimum un nom, un fournisseur, un modele optionnel et un identifiant materiel externe, par exemple IMEI, numero de serie ou identifiant propre au fabricant.")
bullet("Le traceur appartient a un utilisateur AlertContacts, appele proprietaire du traceur.")
bullet("L'onglet Traceurs permet d'ajouter et lister les traceurs. Les APIs backend permettent aussi la modification, la suppression, l'activation, la desactivation, la lecture d'historique et l'association a des zones.")
bullet("Le statut du traceur peut etre draft, active, suspended ou offline.")
bullet("Les positions entrantes sont recues par une route interne serveur a serveur protegee par un secret d'ingestion.")
bullet("Chaque position peut contenir latitude, longitude, precision, vitesse, cap, batterie et date de capture par l'appareil.")
bullet("Le backend met a jour la derniere position, le dernier contact, la batterie et l'historique du traceur.")
bullet("Le geofencing des traceurs traite les zones circulaires actives et cree des evenements entree/sortie avec notification push au proprietaire si la notification est autorisee.")
bullet("Aucun bouton SOS n'est inclus dans cette feature, conformement a la decision produit.")

h2("9.1 Mode gratuit et Premium pour les traceurs")
table(
    ["Fonction", "Gratuit", "Premium"],
    [
        ["Enregistrer un traceur", "Oui, limite a 1 traceur", "Oui, sans limite codee cote API"],
        ["Voir la derniere position connue", "Oui si une position existe deja", "Oui"],
        ["Activer le suivi actif", "Non", "Oui"],
        ["Ingestion de nouvelles positions", "Suspendue si le proprietaire n'est pas Premium", "Oui"],
        ["Historique des positions", "Non", "Oui"],
        ["Association a des zones de securite", "Non", "Oui"],
        ["Alertes entree/sortie de zone", "Non", "Oui"],
    ],
    [3100, 3000, 3260],
    header_fill=LIGHT_BLUE,
)

h2("9.2 Integration fabricant")
paragraph(
    "Le code est volontairement agnostique du fabricant. Tant que le modele exact du traceur n'est pas choisi, AlertContacts doit avancer avec un contrat d'integration stable : le partenaire ou un adaptateur serveur envoie les positions au backend AlertContacts avec provider, external_identifier et coordonnees GPS."
)
bullet("Si le fabricant fournit une API serveur a serveur, l'adaptateur interroge ou recoit les positions puis les transmet a AlertContacts.")
bullet("Si le traceur envoie directement vers une URL, l'URL cible peut etre mappee vers l'endpoint interne d'ingestion, apres authentification par secret ou mecanisme plus robuste.")
bullet("Si le traceur est uniquement BLE/Find My/Find Hub sans GNSS/LTE et sans acces officiel aux positions, il ne repond pas au besoin de suivi autonome.")
bullet("Traccar n'est pas une dependance obligatoire du produit ; il peut rester une option d'integration, mais l'architecture actuelle permet de fonctionner sans lui.")

h1("10. Abonnement et monétisation")
bullet("RevenueCat est la source client des droits Premium, via l'entitlement configure.")
bullet("Le paywall natif RevenueCat UI est presente selon le contexte : limite proche, limite zone, traceurs, mode invisible, etc.")
bullet("Le serveur Laravel reste l'autorite finale pour les routes protegees et les limites d'abonnement.")
bullet("Les prix de reference configures sont 4,99 mensuel et 49,99 annuel avec essai de 14 jours, a confirmer dans Google Play et RevenueCat avant communication commerciale.")
bullet("A l'expiration ou annulation effective d'un abonnement, le webhook RevenueCat repasse l'utilisateur en gratuit et suspend directement les traceurs actifs.")
bullet("La reprise de partage apres mode invisible reste ouverte afin d'eviter qu'un utilisateur reste bloque en etat invisible apres expiration.")

h1("11. Confidentialite, consentement et securite")
bullet("Le partage de position entre proches repose sur une relation acceptee et des permissions de partage.")
bullet("Le mode invisible permet a un utilisateur Premium de suspendre temporairement la publication de sa position.")
bullet("L'application doit eviter tout message laissant penser a une surveillance sans consentement.")
bullet("Les donnees de localisation sont sensibles : elles doivent etre limitees, journalisees avec prudence, securisees en transit et supprimables/exportables selon les demandes RGPD.")
bullet("Les logs, Crashlytics et analytics ne doivent pas exposer d'email, de nom complet ou de coordonnees GPS precises lorsque ce n'est pas strictement necessaire.")
bullet("Les routes internes traceurs doivent rester protegees par un secret fort aujourd'hui, puis evoluer vers une signature HMAC ou un jeton par fournisseur si un fabricant est retenu.")

h1("12. Architecture technique")
table(
    ["Couche", "Technologies et responsabilites"],
    [
        ["Mobile", "Flutter, Provider, GoRouter, Google Maps, Geolocator, Background Fetch, Firebase Auth, FCM, Crashlytics, Analytics, RevenueCat."],
        ["Backend", "Laravel API, Sanctum, MySQL, queues, geoprocessing, webhooks RevenueCat, ingestion traceur."],
        ["Temps reel", "Firebase Realtime Database pour publier les positions mobiles visibles par les proches autorises."],
        ["Notifications", "FCM et notifications locales/critiques selon les cas d'usage."],
        ["Cartographie", "Google Maps cote app, services de lieux/autocompletion et routage cote serveur."],
        ["Abonnement", "RevenueCat cote client, webhook et middleware d'autorisation cote serveur."],
    ],
    [1800, 7560],
)

h1("13. APIs principales")
table(
    ["Domaine", "Endpoints"],
    [
        ["Auth", "POST /api/auth/firebase-login, /login, /register ; POST /api/auth/logout, /refresh ; GET /api/me"],
        ["Utilisateur", "PUT /api/user/profile, POST /api/user/export-data, DELETE /api/user/account"],
        ["Proches", "GET/PUT/DELETE /api/relationships ; invitations ; permissions via PUT /api/contacts/{relationship}/permissions"],
        ["Zones", "GET/POST/PUT/DELETE /api/safe-zones ; assignations ; notification-settings ; API resource /api/zones"],
        ["Localisation mobile", "POST /api/locations/batch, GET /api/locations/recent, POST /api/location/pause, /resume"],
        ["Alertes", "GET/POST /api/alerts ; confirm, deny, report ; API v1 /reports et /incidents"],
        ["Trajets", "POST /api/v1/routes/preview, /avoid, /select, /start, /end, /cancel ; history ; recent-destinations ; avoidance-quota"],
        ["Traceurs", "GET/POST/PUT/DELETE /api/gps-trackers ; activate/deactivate ; locations ; zones ; POST /api/internal/tracker-telemetry"],
        ["Abonnement", "GET /api/subscriptions ; POST /api/webhooks/revenuecat"],
    ],
    [1900, 7460],
)

h1("14. Limites, risques et points a valider")
bullet("Verifier en production que l'entitlement RevenueCat, les produits Google Play, l'essai et les prix sont exactement ceux qui seront communiques.")
bullet("Choisir le fabricant du traceur GPS et obtenir la documentation du protocole, l'identifiant unique, les formats de position, la frequence d'envoi et le modele SIM/eSIM.")
bullet("Renforcer l'authentification de l'ingestion traceur avant usage a grande echelle : secret par fournisseur, signature HMAC, rotation des secrets, anti-rejeu.")
bullet("Confirmer si les zones polygonales doivent etre supportees aussi pour les traceurs, car le service actuel ignore explicitement les polygones.")
bullet("Finaliser les ecrans detail traceur : carte de position, historique visuel, edition, suppression et affectation zone si tout n'est pas encore expose cote Flutter.")
bullet("Unifier les limites d'abonnement entre constantes Flutter, Remote Config eventuel et config Laravel pour eviter les divergences.")
bullet("Verifier les routes profil existantes : le code frontend peut attendre certains endpoints non presents dans les routes actuelles.")
bullet("Maintenir le CDN Hostinger desactive pour le sous-domaine API mobile afin d'eviter les redirections HTML sur les appels JSON.")

h1("15. Criteres d'acceptation")
number("Un compte gratuit peut utiliser l'app sans blocage initial, creer 1 proche, 1 zone et enregistrer 1 traceur GPS.")
number("Au-dela des limites gratuites, le client affiche le paywall et le serveur refuse les actions non autorisees.")
number("Un utilisateur Premium peut activer le mode invisible, ajouter plusieurs proches/zones et utiliser les fonctions avancees de suivi.")
number("Un traceur Premium actif qui envoie une position cree un historique, met a jour sa derniere position et declenche les alertes de zone circulaire si necessaire.")
number("Si l'abonnement Premium expire, les traceurs actifs sont suspendus et l'ingestion de nouvelles positions est refusee.")
number("Les parcours de localisation ne doivent jamais laisser croire a un suivi actif quand le serveur a refuse l'activation.")
number("Les messages marketing et produit doivent parler de rassurance, d'alerte et d'information, jamais de garantie de securite absolue.")

h1("16. Hors perimetre actuel")
bullet("Bouton SOS pour traceur ou application mobile : explicitement non retenu pour cette iteration.")
bullet("Acces aux positions Apple Find My ou Google Find Hub : non garanti sans mecanisme officiel du fabricant et des plateformes.")
bullet("Guidage vocal turn-by-turn complet.")
bullet("Contrat fabricant final, achat de traceurs, gestion de stock et activation SIM/eSIM.")
bullet("Application web publique de suivi des traceurs.")

h1("17. Sources internes consultees")
bullet("Code Flutter : navigation, carte, proches, alertes, trajets, traceurs, abonnements et services de localisation.")
bullet("Code Laravel : routes API, controleurs zones, relations, localisation, incidents, trajets, traceurs, webhooks RevenueCat et middleware Premium.")
bullet("Documents internes : V4_PROGRESS.md, Addendum Trajets V4.1, plan de test V4.1, configuration alertcontacts.php.")

doc.save(OUT)
print(OUT)
